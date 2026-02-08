import os
import glob
import markdown
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
import re


def md_to_docx(md_file, docx_file):
    print(f"Converting {md_file}...")
    try:
        with open(md_file, "r", encoding="utf-8") as f:
            text = f.read()

        # Pre-processing some markdown elements that might not be handled well
        # e.g., removal of complex artifacts links if necessary, or just keeping them as text

        current_md = markdown.markdown(text)
        soup = BeautifulSoup(current_md, "html.parser")

        document = Document()

        # Add Title
        document.add_heading(
            os.path.basename(md_file).replace(".md", "").replace("_", " ").title(), 0
        )

        def add_html_to_doc(soup_elem, doc_elem):
            for child in soup_elem.children:
                if child.name == "h1":
                    doc_elem.add_heading(child.get_text(), level=1)
                elif child.name == "h2":
                    doc_elem.add_heading(child.get_text(), level=2)
                elif child.name == "h3":
                    doc_elem.add_heading(child.get_text(), level=3)
                elif child.name == "h4":
                    doc_elem.add_heading(child.get_text(), level=4)
                elif child.name == "p":
                    # Check for bold/italic in p
                    p = doc_elem.add_paragraph()
                    for content in child.contents:
                        if content.name == "strong" or content.name == "b":
                            p.add_run(content.get_text()).bold = True
                        elif content.name == "em" or content.name == "i":
                            p.add_run(content.get_text()).italic = True
                        elif content.name == "code":
                            run = p.add_run(content.get_text())
                            run.font.name = "Courier New"
                        elif content.name == "a":
                            p.add_run(
                                content.get_text() + f" ({content.get('href', '')})"
                            )
                        elif isinstance(content, str):
                            p.add_run(content)
                elif child.name == "ul":
                    for li in child.find_all("li", recursive=False):
                        doc_elem.add_paragraph(li.get_text(), style="List Bullet")
                elif child.name == "ol":
                    for li in child.find_all("li", recursive=False):
                        doc_elem.add_paragraph(li.get_text(), style="List Number")
                elif child.name == "pre":
                    # Code blocks
                    code = child.get_text()
                    p = doc_elem.add_paragraph(code)
                    p.style = "No Spacing"
                    for run in p.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                elif child.name == "table":
                    # Basic table handling
                    rows = child.find_all("tr")
                    if rows:
                        # Determine max columns
                        max_cols = 0
                        for row in rows:
                            cols = row.find_all(["td", "th"])
                            if len(cols) > max_cols:
                                max_cols = len(cols)

                        table = doc_elem.add_table(rows=len(rows), cols=max_cols)
                        table.style = "Table Grid"

                        for i, row in enumerate(rows):
                            cols = row.find_all(["td", "th"])
                            for j, col in enumerate(cols):
                                table.cell(i, j).text = col.get_text()

        # Iterate over top-level elements
        # BeautifulSoup's direct iteration captures top level tags
        # We need to handle this carefully as soup usually wraps in logic

        # A simpler approach: iterate over all children of body if it existed, or just top level
        # BeautifulSoup automatically adds html/body if not present? No, markdown output is a fragment.

        # Let's just create a wrapper to ensure we have a root
        soup_wrapper = BeautifulSoup(f"<div>{current_md}</div>", "html.parser")
        main_div = soup_wrapper.find("div")

        add_html_to_doc(main_div, document)

        document.save(docx_file)
        print(f"Successfully converted {md_file} to {docx_file}")

    except Exception as e:
        print(f"Error converting {md_file}: {str(e)}")


if __name__ == "__main__":
    files = glob.glob("*.md")
    print(f"Found {len(files)} markdown files.")
    for md_file in files:
        docx_file = md_file.replace(".md", ".docx")
        md_to_docx(md_file, docx_file)
