import shutil
import os
import glob
import markdown
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup


def process_docs():
    # 1. Setup paths
    base_dir = os.getcwd()
    dest_dir = os.path.join(base_dir, "documentations")
    os.makedirs(dest_dir, exist_ok=True)

    # Files to copy
    files_to_copy = [
        "ARCHITECTURE.md",
        "DATABASE.md",
        "README.md",
        "data/data_dictionary.md",
        "data/eda_report.md",
        "docs/admin_guide.md",
        "docs/api_documentation.md",
        "docs/user_manual.md",
        "launch/announcement.md",
        "launch/checklist.md",
        "maintenance/schedule.md",
        "support/procedures.md",
        "tests/README.md",
        "tests/sprint10_testing_guide.md",
        "tests/sprint7_testing_guide.md",
        "tests/sprint8_testing_guide.md",
        "tests/sprint9_testing_guide.md",
        # These were hardcoded to a specific user's brain directory. 
        # We now skip them or use a relative path if they were moved to docs/
        "docs/walkthrough.md",
        "docs/implementation_plan.md",
        "docs/task.md",
    ]

    print("--- STEP 1: Copying Files ---")
    for src_path in files_to_copy:
        try:
            # Handle absolute vs relative
            if not os.path.isabs(src_path):
                full_src = os.path.join(base_dir, src_path)
            else:
                full_src = src_path

            # Normalize for Windows
            full_src = os.path.normpath(full_src)

            if os.path.exists(full_src):
                shutil.copy2(full_src, dest_dir)
                print(f"Copied: {os.path.basename(full_src)}")
            else:
                print(f"Skipped (Not Found): {src_path}")
        except Exception as e:
            print(f"Error copying {src_path}: {e}")

    print("\n--- STEP 2: Converting to DOCX ---")

    md_files = glob.glob(os.path.join(dest_dir, "*.md"))

    for md_file in md_files:
        docx_file = md_file.replace(".md", ".docx")
        print(f"Converting: {os.path.basename(md_file)}")
        try:
            md_to_docx_convert(md_file, docx_file)
        except Exception as e:
            print(f"Failed to convert {os.path.basename(md_file)}: {e}")


def md_to_docx_convert(md_file, docx_file):
    with open(md_file, "r", encoding="utf-8") as f:
        text = f.read()

    # Basic markdown to HTML
    html = markdown.markdown(text, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")

    document = Document()

    # Title from filename
    title = os.path.basename(md_file).replace(".md", "").replace("_", " ").title()
    document.add_heading(title, 0)

    process_element(soup, document)

    document.save(docx_file)


def process_element(soup_elem, doc_elem):
    # Iterate over elements. If soup_elem is the top soup, it has no name, just children
    children = (
        soup_elem.find_all(recursive=False) if hasattr(soup_elem, "find_all") else []
    )
    if not children and hasattr(soup_elem, "contents"):
        # Fallback for when find_all(recursive=False) returns nothing but there are contents (e.g. NavigableString)
        # But usually markdown output is a list of blocks
        pass

    # BeautifulSoup parsing of markdown output usually results in a flat list of tags at top level
    # or wrapped in Data. We iterate over the tags.

    for child in soup_elem:
        if child.name == "h1":
            doc_elem.add_heading(child.get_text(), level=1)
        elif child.name == "h2":
            doc_elem.add_heading(child.get_text(), level=2)
        elif child.name == "h3":
            doc_elem.add_heading(child.get_text(), level=3)
        elif child.name == "h4":
            doc_elem.add_heading(child.get_text(), level=4)
        elif child.name == "p":
            p = doc_elem.add_paragraph()
            add_runs(p, child)
        elif child.name == "ul":
            for li in child.find_all("li", recursive=False):
                p = doc_elem.add_paragraph(style="List Bullet")
                add_runs(p, li)
        elif child.name == "ol":
            for li in child.find_all("li", recursive=False):
                p = doc_elem.add_paragraph(style="List Number")
                add_runs(p, li)
        elif child.name == "pre":
            # Code block
            code = child.get_text()
            p = doc_elem.add_paragraph(code)
            p.style = "No Spacing"  # or generic styling if unavailable
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        elif child.name == "table":
            rows = child.find_all("tr")
            if rows:
                max_cols = 0
                for row in rows:
                    cols = row.find_all(["td", "th"])
                    if len(cols) > max_cols:
                        max_cols = len(cols)

                if max_cols > 0:
                    table = doc_elem.add_table(rows=len(rows), cols=max_cols)
                    table.style = "Table Grid"
                    for i, row in enumerate(rows):
                        cols = row.find_all(["td", "th"])
                        for j, col in enumerate(cols):
                            table.cell(i, j).text = col.get_text()
        elif child.name == "hr":
            doc_elem.add_paragraph(
                "___________________________________________________"
            )


def add_runs(paragraph, soup_element):
    # Recursive function to handle inline formatting (bold, italic, code, link)
    for content in soup_element.contents:
        if content.name == "strong" or content.name == "b":
            run = paragraph.add_run(content.get_text())
            run.bold = True
        elif content.name == "em" or content.name == "i":
            run = paragraph.add_run(content.get_text())
            run.italic = True
        elif content.name == "code":
            run = paragraph.add_run(content.get_text())
            run.font.name = "Courier New"
        elif content.name == "a":
            text = content.get_text()
            href = content.get("href", "")
            run = paragraph.add_run(f"{text} ({href})")
            run.font.color.rgb = None  # Default blue often
            run.underline = True
        elif isinstance(content, str):
            paragraph.add_run(content)
        elif content.name:
            # Recurse for nested tags
            add_runs(paragraph, content)


if __name__ == "__main__":
    process_docs()
